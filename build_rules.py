#!/usr/bin/env python3
"""Генератор правил Чу сёги — A4 DOCX на русском языке."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Цветовая палитра ───────────────────────────────────────────────────────────
C_ACCENT  = RGBColor(0x7C, 0x2D, 0x12)   # бордо — заголовки h1
C_DARK    = RGBColor(0x1F, 0x29, 0x37)   # тёмно-серый — тело
C_MID     = RGBColor(0x6B, 0x5D, 0x4F)   # коричневатый — подписи
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
FILL_HEAD = "F3E8D8"   # кремовый — заголовки таблиц
FILL_ALT  = "FAF7F2"   # светлый — чётные строки
FILL_NONE = None

# ── Шрифты ────────────────────────────────────────────────────────────────────
FONT_RU = "Arial"
FONT_JP = "Yu Gothic"   # fallback: MS Gothic / любой японский шрифт системы

doc = Document()

# ── Размер страницы: A4 ───────────────────────────────────────────────────────
from docx.oxml import parse_xml
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.2)
section.right_margin  = Cm(2.2)

# ── Базовый шрифт документа ───────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = FONT_RU
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = C_DARK
style_normal.paragraph_format.space_after = Pt(5)
style_normal.paragraph_format.line_spacing = Pt(15)

# ── Помощники ─────────────────────────────────────────────────────────────────
def set_cell_shading(cell, fill_hex: str):
    """Залить фоновый цвет ячейки таблицы."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','bottom','left','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '8B7355')
        tcBorders.append(b)
    tcPr.append(tcBorders)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_cell_margins(cell, top=50, bottom=50, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top),('bottom', bottom),('left', left),('right', right)):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def prep_cell(cell, fill=None):
    """Стандартная подготовка ячейки."""
    set_cell_borders(cell)
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))   # 1 cm ≈ 567 twips (DXA)
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def add_run_jp(paragraph, text, size=Pt(13), bold=False, color=C_DARK):
    """Добавить фрагмент с японским шрифтом."""
    run = paragraph.add_run(text)
    run.font.name = FONT_JP
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    # Явно указать East Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_JP)
    return run

def add_run_ru(paragraph, text, size=Pt(11), bold=False, italic=False, color=C_DARK):
    run = paragraph.add_run(text)
    run.font.name = FONT_RU
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run

def heading(level: int, text: str):
    """Добавить заголовок (h1/h2/h3)."""
    sizes  = {1: Pt(18), 2: Pt(14), 3: Pt(12)}
    colors = {1: C_ACCENT, 2: C_DARK, 3: C_DARK}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = add_run_ru(p, text, size=sizes[level], bold=True, color=colors[level])
    return p

def body(text: str, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.5)
    add_run_ru(p, text, italic=italic)
    return p

def bullet_item(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    add_run_ru(p, text)
    return p

def blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

def page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(
        __import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE
    )

# ── Колонтитулы ───────────────────────────────────────────────────────────────
header_sec = doc.sections[0].header
hp = header_sec.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run_jp(hp, '中将棋 ', size=Pt(9))
add_run_ru(hp, '· Чу сёги — правила', size=Pt(9), italic=True, color=C_MID)

footer_sec = doc.sections[0].footer
fp = footer_sec.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run_ru(fp, 'стр. ', size=Pt(9), color=C_MID)
# Номер страницы через поле XML
from docx.oxml import OxmlElement as OE
def add_page_num_field(paragraph):
    run = paragraph.add_run()
    r = run._r
    fldChar = OE('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin'); r.append(fldChar)
    r2 = OE('w:r')
    instrText = OE('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'; r2.append(instrText); paragraph._p.append(r2)
    r3 = OE('w:r')
    fldChar2 = OE('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end'); r3.append(fldChar2)
    paragraph._p.append(r3)

add_page_num_field(fp)
add_run_ru(fp, ' / ', size=Pt(9), color=C_MID)
def add_num_pages_field(paragraph):
    run = paragraph.add_run()
    r = run._r
    fldChar = OE('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin'); r.append(fldChar)
    r2 = OE('w:r')
    instrText = OE('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'NUMPAGES'; r2.append(instrText); paragraph._p.append(r2)
    r3 = OE('w:r')
    fldChar2 = OE('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end'); r3.append(fldChar2)
    paragraph._p.append(r3)

add_num_pages_field(fp)

# ══════════════════════════════════════════════════════════════════════════════
# ОБЛОЖКА
# ══════════════════════════════════════════════════════════════════════════════
cover_ru = doc.add_paragraph()
cover_ru.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_ru.paragraph_format.space_before = Pt(60)
cover_ru.paragraph_format.space_after  = Pt(4)
add_run_ru(cover_ru, 'Чу сёги', size=Pt(32), bold=True, color=C_ACCENT)

cover_jp = doc.add_paragraph()
cover_jp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_jp.paragraph_format.space_after = Pt(8)
add_run_jp(cover_jp, '中将棋', size=Pt(48), bold=True, color=C_DARK)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
add_run_ru(sub, 'Полные правила игры', size=Pt(16), color=C_DARK)

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc.paragraph_format.space_after = Pt(40)
add_run_ru(desc, 'Историческая разновидность японских шахмат · 12×12 · 46 фигур', size=Pt(11), italic=True, color=C_MID)

toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_title.paragraph_format.space_after = Pt(4)
add_run_ru(toc_title, 'Содержание', size=Pt(13), bold=True)

for item in [
    '1.  Введение',
    '2.  Цель игры',
    '3.  Доска и фигуры',
    '4.  Стартовая расстановка',
    '5.  Типы хода',
    '6.  Превращение фигур',
    '7.  Лев и львиная сила',
    '8.  Особые правила взятия Льва',
    '9.  Условия победы',
    '10. Прочие правила',
    '11. Описание всех фигур (справочник)',
]:
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(2)
    add_run_ru(tp, item, size=Pt(11))

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. ВВЕДЕНИЕ
# ══════════════════════════════════════════════════════════════════════════════
heading(1, '1. Введение')
body('Чу сёги («средние сёги») — историческая разновидность японских шахмат, '
     'известная по документам с XIV века и широко игравшаяся в Японии вплоть '
     'до начала XX века. Это прямой предок современных сёги.')
body('Игра считается одной из самых стратегически глубоких в семействе сёги: '
     '144 клетки, 46 фигур у каждого игрока и 21 уникальный тип фигур с '
     'особыми способами хода. В отличие от современных сёги, захваченные '
     'фигуры не возвращаются в игру — это сближает чу сёги с классическими '
     'шахматами и придаёт каждому размену особый вес.')
body('Главная звезда игры — Лев (獅子, шиси). Он настолько силён, что для '
     'него действует отдельный набор правил, ограничивающих его обмен и '
     'определяющих ход игры в целом.')

# ══════════════════════════════════════════════════════════════════════════════
# 2. ЦЕЛЬ ИГРЫ
# ══════════════════════════════════════════════════════════════════════════════
heading(1, '2. Цель игры')
body('Цель — захватить Короля противника. Если у соперника есть Кронпринц '
     '(получается превращением Пьяного слона), нужно захватить обе королевские '
     'фигуры: пока жива хотя бы одна — игрок не проиграл.')
body('Важно: в чу сёги Король именно захватывается, а не «ставится мат» '
     'в формальном смысле. Ходить под бой разрешено. Пата не существует.')

heading(3, 'Альтернативное правило «голого Короля»')
body('По правилам Японской ассоциации чу сёги можно выиграть, оставив '
     'противника с одиноким Королём (голый король), если на следующем ходу '
     'ваш Король не окажется в таком же положении. Договоритесь с соперником '
     'до партии, используете ли это правило.')

# ══════════════════════════════════════════════════════════════════════════════
# 3. ДОСКА И ФИГУРЫ
# ══════════════════════════════════════════════════════════════════════════════
heading(1, '3. Доска и фигуры')
body('Доска квадратная, 12 рядов × 12 файлов = 144 клетки. Клетки не '
     'раскрашены. Принадлежность фигуры определяется направлением её '
     'заострённого конца: остриё смотрит на противника.')
body('У каждого игрока 46 фигур 21 типа. Имена написаны иероглифами кандзи. '
     'Превращение обозначается переворачиванием фигуры на оборотную сторону.')

heading(3, 'Состав армии (по 46 у каждого игрока)')

army = [
    ('Король (王将 / 玉将)', '1'),
    ('Свободный король / Ферзь (奔王)', '1'),
    ('Лев (獅子)', '1'),
    ('Пьяный слон (酔象)', '1'),
    ('Кирин (麒麟)', '1'),
    ('Феникс (鳳凰)', '1'),
    ('Дракон-король (龍王)', '2'),
    ('Дракон-конь (龍馬)', '2'),
    ('Ладья (飛車)', '2'),
    ('Слон (角行)', '2'),
    ('Вертикальный ходок (竪行)', '2'),
    ('Боковой ходок (横行)', '2'),
    ('Колесница назад (反車)', '2'),
    ('Копьё (香車)', '2'),
    ('Слепой тигр (盲虎)', '2'),
    ('Золотой генерал (金将)', '2'),
    ('Серебряный генерал (銀将)', '2'),
    ('Медный генерал (銅将)', '2'),
    ('Свирепый барс (猛豹)', '2'),
    ('Посредник (仲人)', '2'),
    ('Пешка (歩兵)', '12'),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.autofit = False
# Widths: 13 cm + 3 cm = 16 cm content
W_ARMY_NAME = int(13 * 567)
W_ARMY_QTY  = int(3  * 567)

hrow = tbl.rows[0]
for cell, (txt_h, w) in zip(hrow.cells, [('Фигура', W_ARMY_NAME), ('Кол-во', W_ARMY_QTY)]):
    prep_cell(cell, FILL_HEAD)
    set_col_width(cell, w / 567)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_ru(p, txt_h, size=Pt(10), bold=True)

for i, (name, qty) in enumerate(army):
    row = tbl.add_row()
    fill = FILL_ALT if i % 2 else FILL_NONE
    c0, c1 = row.cells
    prep_cell(c0, fill); set_col_width(c0, W_ARMY_NAME / 567)
    prep_cell(c1, fill); set_col_width(c1, W_ARMY_QTY / 567)
    add_run_ru(c0.paragraphs[0], name, size=Pt(10))
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_ru(c1.paragraphs[0], qty, size=Pt(10))

blank()

# ══════════════════════════════════════════════════════════════════════════════
# 4. СТАРТОВАЯ РАССТАНОВКА
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading(1, '4. Стартовая расстановка')
body('Каждая сторона занимает 4 ряда: три ряда фигур у тыла + ряд из 12 пешек, '
     'прикрытых двумя посредниками спереди.')
body('Расстановка асимметрична: пары «Король / Пьяный слон», «Лев / Свободный '
     'король» у белых зеркально переставлены относительно чёрных. Король '
     'чёрных стоит напротив Пьяного слона белых — это фирменная черта чу сёги.')

heading(3, 'Расстановка чёрных (тыл → фронт, файлы 1–12)')

# Setup table: 13 rows (label) + 12 kanji cols
setup_rows = [
    ('Ряд 1 (тыл)',    ['香','豹','銅','銀','金','王','象','金','銀','銅','豹','香']),
    ('Ряд 2',          ['反','　','角','　','虎','麒','鳳','虎','　','角','　','反']),
    ('Ряд 3',          ['横','竪','飛','馬','龍','獅','奔','龍','馬','飛','竪','横']),
    ('Ряд 4 (пешки)',  ['歩','歩','歩','歩','歩','歩','歩','歩','歩','歩','歩','歩']),
    ('Ряд 5 (посред.)',['　','　','　','仲','　','　','　','　','仲','　','　','　']),
]
COL_LABEL = 3.8
COL_FILE  = 1.0  # 12 × 1.0 = 12 cm

stbl = doc.add_table(rows=1 + len(setup_rows), cols=13)
stbl.style = 'Table Grid'
stbl.alignment = WD_TABLE_ALIGNMENT.CENTER
stbl.autofit = False

# Header row
hdr = stbl.rows[0].cells
prep_cell(hdr[0], FILL_HEAD); set_col_width(hdr[0], COL_LABEL)
p_h = hdr[0].paragraphs[0]
add_run_ru(p_h, 'Ряд  \\ файл', size=Pt(8), bold=True)

for i, num in enumerate('1 2 3 4 5 6 7 8 9 10 11 12'.split()):
    c = hdr[i + 1]
    prep_cell(c, FILL_HEAD); set_col_width(c, COL_FILE)
    p_c = c.paragraphs[0]
    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_ru(p_c, num, size=Pt(8), bold=True)

for ri, (label, glyphs) in enumerate(setup_rows):
    row = stbl.rows[ri + 1].cells
    prep_cell(row[0])
    add_run_ru(row[0].paragraphs[0], label, size=Pt(8))
    set_col_width(row[0], COL_LABEL)
    for ci, g in enumerate(glyphs):
        c = row[ci + 1]
        prep_cell(c, FILL_ALT if ri % 2 else FILL_NONE)
        set_col_width(c, COL_FILE)
        p_c = c.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if g.strip():
            add_run_jp(p_c, g, size=Pt(12), bold=True)

blank()

heading(3, 'Ключ кандзи')
key_data = [
    ('香','Копьё'),('豹','Свирепый барс'),('銅','Медный'),('銀','Серебряный'),
    ('金','Золотой генерал'),('王','Король'),('象','Пьяный слон'),
    ('反','Колесница назад'),('角','Слон'),('虎','Слепой тигр'),
    ('麒','Кирин'),('鳳','Феникс'),('横','Боков. ходок'),('竪','Верт. ходок'),
    ('飛','Ладья'),('馬','Дракон-конь'),('龍','Дракон-король'),
    ('獅','Лев'),('奔','Свободный король'),('歩','Пешка'),('仲','Посредник'),
]
# 3-колонная таблица-ключ
def chunks(lst, n): return [lst[i:i+n] for i in range(0, len(lst), n)]
key_chunks = chunks(key_data, 3)
ktbl = doc.add_table(rows=len(key_chunks), cols=3)
ktbl.style = 'Table Grid'
ktbl.alignment = WD_TABLE_ALIGNMENT.CENTER
ktbl.autofit = False
for ri, row_data in enumerate(key_chunks):
    row = ktbl.rows[ri].cells
    for ci in range(3):
        prep_cell(row[ci])
        set_col_width(row[ci], 16.0 / 3 / 10)  # ~5.3 cm each? → total 16cm
        if ci < len(row_data):
            k, name = row_data[ci]
            p = row[ci].paragraphs[0]
            add_run_jp(p, k + ' ', size=Pt(13), bold=True, color=C_ACCENT)
            add_run_ru(p, '— ' + name, size=Pt(9))

blank()
body('Расстановка белых зеркальна, но Король белых в файле 7, а не 6 '
     '(а Пьяный слон — в файле 6). Первый ход делают чёрные.')

# ══════════════════════════════════════════════════════════════════════════════
# 5. ТИПЫ ХОДА
# ══════════════════════════════════════════════════════════════════════════════
heading(1, '5. Типы хода')

heading(3, 'Шаговый ход')
body('Фигура перемещается ровно на одну клетку. Свою фигуру не обойти; '
     'чужую можно взять.')
body('Шаговые фигуры: Король, Пьяный слон, Слепой тигр, Свирепый барс, '
     'все генералы (Золотой / Серебряный / Медный), Посредник, Пешка.')

heading(3, 'Дальнобойный ход')
body('Фигура движется на любое число свободных клеток вдоль одной линии. '
     'Чужую фигуру на пути можно взять (остановившись на её клетке); '
     'свою — нельзя миновать. Дальнобойные: Свободный король (Ферзь), '
     'Ладья, Слон, Дракон-король, Дракон-конь, а также Боковой / '
     'Вертикальный ходок, Колесница назад, Копьё.')

heading(3, 'Прыжковый ход')
body('Фигура прыгает через одну клетку, игнорируя любую фигуру на '
     'промежуточной клетке. Прыжковые: Лев, Кирин, Феникс, '
     'а также превращённые Дракон-конь / Дракон-король.')

heading(3, 'Львиная сила')
body('Особый ход: два шага подряд за один ход (см. раздел 7).')

# ══════════════════════════════════════════════════════════════════════════════
# 6. ПРЕВРАЩЕНИЕ
# ══════════════════════════════════════════════════════════════════════════════
heading(1, '6. Превращение фигур')
body('Зона превращения — четыре ряда наиболее близкие к противнику. '
     'Для чёрных это ряды 9–12, для белых — ряды 1–4.')
body('Превращение ДОБРОВОЛЬНО (если хотите, превращайте; нет — подождите). '
     'Совет: превращайте при первой возможности — большинство превращений '
     'дают значительное усиление.')

heading(3, 'Когда превращение разрешено')
bullet_item('Фигура ВПЕРВЫЕ ВХОДИТ в зону превращения (из-за её пределов).')
bullet_item('Фигура БЕРЁТ фигуру противника, и хотя бы часть хода '
            'затрагивает зону или происходит при выходе из зоны.')

heading(3, 'Когда превращение НЕ разрешено')
bullet_item('Фигура уже стоит в зоне и просто переставляется внутри неё '
            'без взятия — превратиться нельзя.')
bullet_item('Если вы вошли в зону и не превратились, позже можно '
            'превратиться только: (а) выйдя из зоны и снова войдя, '
            'либо (б) совершив взятие в зоне.')
body('Примечание: это строже, чем в современных сёги.')

heading(3, 'Особые случаи')
bullet_item('Пешка: при достижении последнего ряда без взятия всё равно '
            'может (и должна) превратиться, чтобы не застрять.')
bullet_item('Копьё: если не превратилось у последнего ряда — навсегда '
            'теряет подвижность.')
bullet_item('Король, Лев, Свободный король — не превращаются.')
bullet_item('Пьяный слон превращается в Кронпринца (太子) — вторую '
            'королевскую фигуру, двигающуюся как Король.')

# ══════════════════════════════════════════════════════════════════════════════
# 7. ЛЕВ
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading(1, '7. Лев и львиная сила')
body('Лев (獅子) — самая сильная фигура. За один ход он может охватить '
     'все 24 клетки в радиусе двух шагов от себя.')

heading(3, 'Способы хода Льва')
bullet_item('Прыжок на 2 клетки в любом из 8 направлений (игнорируя '
            'промежуточную клетку).')
bullet_item('Два последовательных шага в произвольных направлениях '
            'с возможным взятием на каждом шаге — Лев может забрать '
            'сразу две фигуры за один ход.')
bullet_item('Игуй (居食い): перейти на соседнюю клетку, взять там фигуру '
            'и вернуться — взятие «не сходя с места».')
bullet_item('Дзитто (じっと): перейти на соседнюю клетку и вернуться '
            'без взятия — фактически пропуск хода.')

heading(3, 'Льва также имеют')
bullet_item('Сокол с рогами (角鷹) — превращённый Дракон-конь: '
            'ограниченная львиная сила в двух клетках строго вперёд.')
bullet_item('Парящий орёл (飛鷲) — превращённый Дракон-король: '
            'ограниченная львиная сила по двум диагоналям вперёд.')

# ══════════════════════════════════════════════════════════════════════════════
# 8. ПРАВИЛА ВЗЯТИЯ ЛЬВА
# ══════════════════════════════════════════════════════════════════════════════
heading(1, '8. Особые правила взятия Льва')
body('Лев слишком силён для простого обмена, поэтому существует '
     'набор правил, не дающих быстро «разменять» Льва.')

heading(3, 'Главный принцип')
body('Если ваш Лев только что захвачен, на следующем ходу вы '
     'НЕ МОЖЕТЕ взять вражеского Льва не-Львом. Нужно сначала '
     'пропустить ход или использовать своего Льва.')

heading(3, 'Не-Лев может взять Льва, если:')
bullet_item('Лев противника не защищён ни одной его фигурой '
            '(никто не может сразу перебить).')
bullet_item('Лев защищён только Пешкой или Посредником — такая '
            'защита «не считается».')
bullet_item('Не-Лев забирает Льва И дополнительно ещё одну ценную '
            'фигуру (не Пешку, не Посредника) с помощью двойного хода '
            'Сокола или Орла — ценная добыча оправдывает риск.')
bullet_item('Взятие является ответным ударом: противник на прошлом ходу '
            'сам взял вашу фигуру не-Львом.')

heading(3, 'Правило «мост» между Львами')
body('Если между двумя Львами стоит фигура противника (не Пешка / '
     'Посредник), ситуация усложняется: Лев может взять «мост» '
     'первым шагом и Льва вторым — тогда соперник может ответить. '
     'Согласуйте с партнёром до игры подробный вариант этого правила '
     'или пользуйтесь упрощённой версией (ниже).')

heading(3, 'Упрощённая версия для первой партии')
body('«Не-Лев может взять Льва, только если тот не защищён». '
     'Этого достаточно для игры в начале освоения чу сёги.')

# ══════════════════════════════════════════════════════════════════════════════
# 9. УСЛОВИЯ ПОБЕДЫ
# ══════════════════════════════════════════════════════════════════════════════
heading(1, '9. Условия победы')
body('Победа — захват всех королевских фигур противника. '
     'Это может быть один Король или Король + Кронпринц.')
body('Правило «голого Короля» (опционально): игра заканчивается победой, '
     'если у противника не осталось ни одной некоролевской фигуры, '
     'при условии что ваш Король сам не «оголяется» на следующем ходу. '
     'Если же и ваш Король оголяется — ничья.')
body('Исключение: если у выигрывающей стороны остались только '
     'Король + Пешка или Король + Посредник, для победы по этому '
     'правилу пешка или посредник должны сначала превратиться.')

# ══════════════════════════════════════════════════════════════════════════════
# 10. ПРОЧИЕ ПРАВИЛА
# ══════════════════════════════════════════════════════════════════════════════
heading(1, '10. Прочие правила')

heading(3, 'Запрет повторения')
body('Нельзя совершить ход, после которого позиция (с тем же ходящим) '
     'уже встречалась. Тот, кто создаёт повторение, обязан его избежать; '
     'при вечном шахе проигрывает дающий шах.')

heading(3, 'Нет сбросов')
body('Захваченные фигуры удаляются с доски навсегда — их нельзя '
     'вернуть в игру, как в современных сёги.')

heading(3, 'Незаконный ход')
body('Совершение незаконного хода ведёт к немедленному проигрышу. '
     'В дружеских партиях обычно разрешают взять ход назад.')

heading(3, 'Двойной запрет (аналог нифу)')
body('Нельзя иметь две непревращённые Пешки в одном файле '
     '(столбце). Аналогично для Посредника: два непревращённых '
     'Посредника в одном файле запрещены.')

heading(3, 'Гандикап')
body('Игры между неравными по силе партнёрами часто ведутся с форой: '
     'у белых убирают одну или несколько фигур, взамен белые ходят первыми '
     'и могут переставить оставшиеся фигуры.')

# ══════════════════════════════════════════════════════════════════════════════
# 11. СПРАВОЧНИК ФИГУР
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading(1, '11. Описание всех фигур (справочник)')
body('Все 21 тип в порядке: сначала шаговые, затем прыжковые, '
     'затем дальнобойные, и наконец Лев. '
     'Сокращение кандзи — то, что нанесено на саму фигуру.')

# Columns: kanji (2.2cm) | name (3.5cm) | qty (1cm) | move (6.5cm) | promo (4.5cm) = 17.7cm
COL_W = [2.2, 3.5, 1.0, 6.5, 4.5]  # cm

pieces = [
    # (full_kanji, short_kanji, ru_name, translit, qty, move_desc, promo_to)
    # ── ШАГОВЫЕ ──────────────────────────────────────────────────────────────
    ('歩兵','歩','Пешка','fuhyō','12',
     'Один шаг строго вперёд.',
     'Золотой генерал (と)'),
    ('仲人','仲','Посредник','chūnin','2',
     'Один шаг строго вперёд или строго назад.',
     'Пьяный слон'),
    ('銅将','銅','Медный генерал','dōshō','2',
     'Один шаг вперёд, по диагоналям вперёд или назад. 5 направлений.',
     'Боковой ходок'),
    ('銀将','銀','Серебряный генерал','ginshō','2',
     'Один шаг вперёд или по любой из 4 диагоналей. 5 направлений.',
     'Вертикальный ходок'),
    ('金将','金','Золотой генерал','kinshō','2',
     'Один шаг в любом направлении, кроме двух диагоналей назад. 6 направлений.',
     'Ладья'),
    ('猛豹','豹','Свирепый барс','mōhyō','2',
     'Один шаг в любом направлении, кроме строго влево/вправо. 6 направлений.',
     'Слон'),
    ('盲虎','虎','Слепой тигр','mōko','2',
     'Один шаг в любом направлении, кроме строго вперёд. 7 направлений.',
     'Свободный король'),
    ('酔象','象','Пьяный слон','suizō','1',
     'Один шаг в любом направлении, кроме строго назад. 7 направлений.',
     'Кронпринц (太子) — как Король; королевская фигура'),
    ('王将 / 玉将','王 / 玉','Король','ōshō','1',
     'Один шаг в любом из 8 направлений.',
     'Не превращается'),
    # ── ПРЫЖКОВЫЕ ────────────────────────────────────────────────────────────
    ('麒麟','麒','Кирин','kirin','1',
     'Один шаг по любой диагонали ИЛИ прыжок на 2 клетки прямо/назад/влево/вправо.',
     'Лев (獅子)'),
    ('鳳凰','鳳','Феникс','hōō','1',
     'Один шаг прямо/назад/влево/вправо ИЛИ прыжок на 2 клетки по любой диагонали.',
     'Свободный король'),
    # ── ДАЛЬНОБОЙНЫЕ ─────────────────────────────────────────────────────────
    ('香車','香','Копьё','kyōsha','2',
     'Любое число клеток строго вперёд.',
     'Белая лошадь (白駒) — слон + дальнобойно вперёд/назад'),
    ('反車','反','Колесница назад','hansha','2',
     'Любое число клеток строго вперёд или строго назад.',
     'Дракон-конь'),
    ('横行','横','Боковой ходок','ōgyō','2',
     'Любое число клеток влево/вправо + один шаг вперёд или назад.',
     'Свободная колесница — то же + дальнобойно вперёд/назад'),
    ('竪行','竪','Вертикальный ходок','shugyō','2',
     'Любое число клеток вперёд/назад + один шаг влево или вправо.',
     'Летящий бык — то же + дальнобойно по диагоналям'),
    ('角行','角','Слон','kakugyō','2',
     'Любое число клеток по любой из 4 диагоналей.',
     'Дракон-конь (龍馬)'),
    ('飛車','飛','Ладья','hisha','2',
     'Любое число клеток прямо / назад / влево / вправо.',
     'Дракон-король (龍王)'),
    ('龍馬','馬','Дракон-конь','ryūma','2',
     'Дальнобойно по диагоналям + один шаг прямо/назад/влево/вправо.',
     'Сокол с рогами (角鷹) — то же + львиная сила вперёд'),
    ('龍王','龍','Дракон-король','ryūō','2',
     'Дальнобойно прямо/назад/влево/вправо + один шаг по любой диагонали.',
     'Парящий орёл (飛鷲) — то же + львиная сила по диагоналям вперёд'),
    ('奔王','奔','Свободный король','honnō','1',
     'Любое число клеток в любом из 8 направлений (как ферзь в шахматах).',
     'Не превращается'),
    # ── ЛЕВ ──────────────────────────────────────────────────────────────────
    ('獅子','獅','Лев','shishi','1',
     'Двойной ход (два шага в произвольных направлениях) или прыжок на 2 клетки. '
     'Может взять до двух фигур за ход. Игуй, Дзитто. См. раздел 7.',
     'Не превращается'),
]

ptbl = doc.add_table(rows=1, cols=5)
ptbl.style = 'Table Grid'
ptbl.alignment = WD_TABLE_ALIGNMENT.CENTER
ptbl.autofit = False

heads = ['Кандзи', 'Название', 'Кол.', 'Ход', 'Превращается в']
for ci, (hd, w) in enumerate(zip(heads, COL_W)):
    c = ptbl.rows[0].cells[ci]
    prep_cell(c, FILL_HEAD)
    set_col_width(c, w)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_ru(p, hd, size=Pt(9), bold=True)

for ri, (full_k, short_k, ru_name, translit, qty, move_d, promo) in enumerate(pieces):
    fill = FILL_ALT if ri % 2 else FILL_NONE
    row = ptbl.add_row()
    for ci, w in enumerate(COL_W):
        prep_cell(row.cells[ci], fill)
        set_col_width(row.cells[ci], w)

    # Kanji col
    p0 = row.cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_jp(p0, full_k, size=Pt(10), bold=True)
    p0.add_run('\n')
    add_run_jp(p0, short_k, size=Pt(13), bold=True, color=C_ACCENT)

    # Name col
    p1 = row.cells[1].paragraphs[0]
    add_run_ru(p1, ru_name, size=Pt(9), bold=True)
    p1.add_run('\n')
    add_run_ru(p1, translit, size=Pt(8), italic=True, color=C_MID)

    # Qty
    p2 = row.cells[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_ru(p2, qty, size=Pt(9))

    # Move
    add_run_ru(row.cells[3].paragraphs[0], move_d, size=Pt(9))

    # Promo
    add_run_ru(row.cells[4].paragraphs[0], promo, size=Pt(9))

blank()
heading(3, 'Краткая сетка превращений')
for line in [
    'Пешка → Золотой генерал (то-кин)',
    'Медный генерал → Боковой ходок',
    'Серебряный генерал → Вертикальный ходок',
    'Золотой генерал → Ладья',
    'Свирепый барс → Слон',
    'Слепой тигр → Свободный король',
    'Феникс → Свободный король',
    'Посредник → Пьяный слон',
    'Копьё → Белая лошадь',
    'Колесница назад / Слон → Дракон-конь',
    'Ладья → Дракон-король',
    'Дракон-конь → Сокол с рогами (+ львиная сила вперёд)',
    'Дракон-король → Парящий орёл (+ львиная сила по диагоналям)',
    'Кирин → Лев  ⚠ (второй Лев — мощно, но рискованно)',
    'Пьяный слон → Кронпринц (вторая королевская фигура)',
    'Король / Лев / Свободный король — не превращаются',
]:
    bullet_item(line)

blank()
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run_ru(p_end, 'Удачной игры!', size=Pt(14), bold=True, color=C_ACCENT)

# ── Сохранить ─────────────────────────────────────────────────────────────────
import os; os.makedirs("ru", exist_ok=True)
out = "ru/chu_shogi_rules.docx"
doc.save(out)
print(f'Saved: {out}')
